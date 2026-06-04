"""Shared pytest fixtures and configuration.

Each unit test should be hermetic: no real API calls, no real disk writes
outside `tmp_path`. We enforce that by clearing the settings cache between
tests (so env-var changes inside a test take effect) and by skipping the
integration suite unless `-m integration` is passed.

Two autouse fixtures live here:

* `_reset_settings_cache` (function-scoped) — clears the memoized
  `get_settings()` between tests so env-var changes take effect.
* `_generate_sample_fixtures` (session-scoped) — synthesizes PDF and
  DOCX fixtures once per session. Markdown / text fixtures are
  hand-written and committed; binary formats are generated on the fly
  so they don't bloat the repo.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from src.config import settings as settings_module


# ─── Settings cache reset (function-scoped, autouse) ─────────────────────


@pytest.fixture(autouse=True)
def _reset_settings_cache() -> None:
    """Clear the memoized `get_settings()` so each test sees fresh env vars."""
    settings_module.get_settings.cache_clear()
    yield
    settings_module.get_settings.cache_clear()


# ─── Sample-document fixture generation (session-scoped, autouse) ────────


_FIXTURES_DIR = Path(__file__).parent / "evals" / "fixtures" / "sample_docs"


@pytest.fixture(scope="session", autouse=True)
def _generate_sample_fixtures() -> None:
    """Synthesize binary sample fixtures once per test session.

    Creates three files inside ``tests/evals/fixtures/sample_docs/`` if
    they don't already exist:

    * ``sample_pdf_simple.pdf``     — 2 pages, hand-typed text
    * ``sample_pdf_multipage.pdf``  — 5 pages, hand-typed text
    * ``sample_document.docx``      — multiple paragraphs + 2 headings

    The Markdown / text fixtures are checked in directly (see
    `tests/evals/fixtures/sample_docs/`); only PDF and DOCX are
    generated because committing binary blobs makes diffs and merges
    miserable.

    If the optional deps aren't installed (rare — they're listed in
    `pyproject.toml`), generation is silently skipped and the
    individual loader tests that need binary fixtures will themselves
    skip via `pytest.importorskip`.
    """
    _FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    _generate_simple_pdf(_FIXTURES_DIR / "sample_pdf_simple.pdf")
    _generate_multipage_pdf(_FIXTURES_DIR / "sample_pdf_multipage.pdf")
    _generate_docx(_FIXTURES_DIR / "sample_document.docx")
    _generate_scanned_pdf(_FIXTURES_DIR / "sample_pdf_scanned.pdf")
    yield


def _generate_simple_pdf(path: Path) -> None:
    """2-page PDF — used to validate basic per-page extraction."""
    if path.exists():
        return
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return  # `pytest.importorskip` in the test gates downstream skips.

    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawString(50, 750, "Sample PDF - Page 1")
    c.drawString(50, 720, "This is a test PDF document with simple text.")
    c.drawString(50, 690, "Used for testing PDFLoader implementation.")
    c.showPage()

    c.drawString(50, 750, "Sample PDF - Page 2")
    c.drawString(50, 720, "This is the second page of the test PDF.")
    c.drawString(50, 690, "It verifies multi-page extraction works.")
    c.showPage()
    c.save()


def _generate_multipage_pdf(path: Path) -> None:
    """5-page PDF — used to validate page-order preservation."""
    if path.exists():
        return
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return

    c = canvas.Canvas(str(path), pagesize=letter)
    for page_num in range(1, 6):
        c.drawString(50, 750, f"Sample PDF - Page {page_num}")
        c.drawString(50, 720, f"This is page {page_num} of a multi-page document.")
        c.drawString(50, 690, f"Content for testing page extraction: page {page_num}.")
        c.showPage()
    c.save()


def _generate_docx(path: Path) -> None:
    """DOCX with multiple paragraphs + a heading — DocxLoader fixture."""
    if path.exists():
        return
    try:
        from docx import Document
    except ImportError:
        return

    doc = Document()
    doc.add_heading("Sample DOCX Document", 0)
    doc.add_paragraph(
        "This is a test DOCX file used for testing the DocxLoader implementation. "
        "It contains multiple paragraphs to verify paragraph extraction works correctly."
    )
    doc.add_paragraph(
        "Here is the second paragraph. It has different content to test that "
        "multiple paragraphs are properly extracted and separated."
    )
    doc.add_heading("Section 2", level=1)
    doc.add_paragraph(
        "This paragraph is under Section 2. It tests header handling and "
        "paragraph ordering in DOCX files."
    )
    doc.save(str(path))


# Common system-font paths we try for the synthetic scanned PDF, in
# rough preference order. We fall back to PIL's bitmap default if none
# resolves — the OCR test that depends on this fixture will simply skip
# rather than fail in that case.
_SCANNED_FONT_CANDIDATES: tuple[str, ...] = (
    "/System/Library/Fonts/Helvetica.ttc",                       # macOS
    "/System/Library/Fonts/Supplemental/Arial.ttf",              # macOS
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",      # Debian / Ubuntu
    "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",               # Fedora / Arch
    "C:/Windows/Fonts/Arial.ttf",                                # Windows
)


def _generate_scanned_pdf(path: Path) -> None:
    """Image-only single-page PDF — no extractable text layer.

    Used by the OCR tests to exercise `PDFLoader(enable_ocr_fallback=True)`
    and `ScannedPDFLoader` without needing a real scanned document
    committed to the repo. We render the text into a PIL image, save
    the image into a PDF page via reportlab, and don't draw any text
    objects directly — so `pypdf.extract_text()` returns essentially
    nothing and only OCR can recover the content.
    """
    if path.exists():
        return
    try:
        import io

        from PIL import Image, ImageDraw, ImageFont
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.utils import ImageReader
        from reportlab.pdfgen import canvas
    except ImportError:
        return

    font = None
    for fp in _SCANNED_FONT_CANDIDATES:
        try:
            font = ImageFont.truetype(fp, 32)
            break
        except (OSError, IOError):
            continue
    if font is None:
        # The default bitmap font is small (~11px) and OCRs unreliably.
        # We still emit the PDF; the OCR test will skip cleanly if the
        # tesseract pass yields nothing meaningful.
        font = ImageFont.load_default()

    # 2x letter at ~150 DPI — high enough that tesseract has decent
    # signal even on the bitmap fallback font.
    img = Image.new("RGB", (1224, 1584), "white")
    draw = ImageDraw.Draw(img)
    for line_idx, line in enumerate(
        [
            "Sample Scanned PDF",
            "This page has no text layer.",
            "OCR is required to extract this content.",
            "Hello from the scanned fixture.",
        ]
    ):
        draw.text((80, 80 + line_idx * 80), line, fill="black", font=font)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)

    c = canvas.Canvas(str(path), pagesize=letter)
    c.drawImage(ImageReader(buf), 0, 0, width=612, height=792)
    c.showPage()
    c.save()
